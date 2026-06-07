#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

DOCX_IMPORT_ERROR = None

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
except ImportError as exc:
    DOCX_IMPORT_ERROR = exc
    Document = None
    WD_ALIGN_VERTICAL = None
    WD_ROW_HEIGHT_RULE = None
    DocxTable = None
    DocxParagraph = None

    class _DummyAlign:
        LEFT = 0
        CENTER = 1
        RIGHT = 2
        JUSTIFY = 3

    class _DummyBreak:
        PAGE = 7

    class _DummyLineSpacing:
        SINGLE = 0
        ONE_POINT_FIVE = 1

    WD_ALIGN_PARAGRAPH = _DummyAlign
    WD_BREAK = _DummyBreak
    WD_LINE_SPACING = _DummyLineSpacing

    def OxmlElement(*args, **kwargs):
        raise RuntimeError('python-docx is required to build DOCX output')

    def qn(name):
        return name

    def Cm(value):
        return value

    def Pt(value):
        return value

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_MONO = "Menlo"
MAX_HEADING_LEVEL = 9
SIZE_TITLE = Pt(28)
SIZE_H1 = Pt(16)
SIZE_H2 = Pt(15)
SIZE_H3 = Pt(14)
SIZE_H4 = Pt(12)
SIZE_H5 = Pt(12)
SIZE_H6 = Pt(12)
SIZE_H7 = Pt(12)
SIZE_H8 = Pt(12)
SIZE_H9 = Pt(12)
SIZE_BODY = Pt(12)
SIZE_TABLE = Pt(10.5)
SIZE_CODE = Pt(9)
A4_WIDTH = Cm(21)
A4_HEIGHT = Cm(29.7)
PAGE_MARGIN_TOP = Cm(2.54)
PAGE_MARGIN_BOTTOM = Cm(2.54)
PAGE_MARGIN_LEFT = Cm(3.17)
PAGE_MARGIN_RIGHT = Cm(3.17)
TABLE_CONTENT_WIDTH_DXA = 11906 - 1800 - 1800
TABLE_CELL_MARGIN_TOP = 80
TABLE_CELL_MARGIN_BOTTOM = 80
TABLE_CELL_MARGIN_LEFT = 120
TABLE_CELL_MARGIN_RIGHT = 120
CODE_LEFT_INDENT = Cm(0.74)
CODE_RIGHT_INDENT = Cm(0.74)
CODE_SHADE = "F2F2F2"
TABLE_HEADER_SHADE = "D9EAF7"
COVER_TOP_SPACER_COUNT = 6
TOC_TITLE = "目录"
TOC_HEADING_MARKERS = {"目录", "目 录", "toc", "table of contents", "contents"}
TXT_HEADING_MAX_LENGTH = 40
SUPPORTED_INPUT_SUFFIXES = {'.md', '.markdown', '.txt', '.docx'}
TITLE_NUMBERING_ABSTRACT_ID = '700'
TITLE_NUMBERING_ID = '701'
BULLET_NUMBERING_ABSTRACT_ID = '702'
BULLET_NUMBERING_ID = '703'


class StyleSpec:
    def __init__(self, font_name: str, font_size, line_spacing: str, before_pt: float, after_pt: float, first_line_chars: int = 0, bold: bool | None = None, alignment=None, font_color: str | None = None):
        self.font_name = font_name
        self.font_size = font_size
        self.line_spacing = line_spacing
        self.before_pt = before_pt
        self.after_pt = after_pt
        self.first_line_chars = first_line_chars
        self.bold = bold
        self.alignment = alignment if alignment is not None else WD_ALIGN_PARAGRAPH.LEFT
        self.font_color = font_color


BLACK_FONT_COLOR = '000000'


TITLE_SPEC = StyleSpec(FONT_HEI, SIZE_TITLE, "single", 10, 10, 0, True, font_color=BLACK_FONT_COLOR)
H1_SPEC = StyleSpec(FONT_HEI, SIZE_H1, "single", 10, 10, 0, True, font_color=BLACK_FONT_COLOR)
H2_SPEC = StyleSpec(FONT_HEI, SIZE_H2, "single", 8, 8, 0, True, font_color=BLACK_FONT_COLOR)
H3_SPEC = StyleSpec(FONT_HEI, SIZE_H3, "one_point_five", 0, 0, 0, True, font_color=BLACK_FONT_COLOR)
H4_SPEC = StyleSpec(FONT_HEI, SIZE_H4, "one_point_five", 0, 0, 0, True, font_color=BLACK_FONT_COLOR)
H5_SPEC = StyleSpec(FONT_HEI, SIZE_H5, "one_point_five", 0, 0, 0, True, font_color=BLACK_FONT_COLOR)
H6_SPEC = StyleSpec(FONT_HEI, SIZE_H6, "one_point_five", 0, 0, 0, True, font_color=BLACK_FONT_COLOR)
H7_SPEC = StyleSpec(FONT_HEI, SIZE_H7, "one_point_five", 0, 0, 0, True, font_color=BLACK_FONT_COLOR)
H8_SPEC = StyleSpec(FONT_HEI, SIZE_H8, "one_point_five", 0, 0, 0, True, font_color=BLACK_FONT_COLOR)
H9_SPEC = StyleSpec(FONT_HEI, SIZE_H9, "one_point_five", 0, 0, 0, True, font_color=BLACK_FONT_COLOR)
COVER_TITLE_SPEC = StyleSpec(FONT_SONG, Pt(22), "one_point_five", 9, 9, 0, True, font_color=BLACK_FONT_COLOR)
COVER_META_SPEC = StyleSpec(FONT_SONG, Pt(14), "one_point_five", 0, 0, 0, None)
COVER_CORNER_SPEC = StyleSpec(FONT_HEI, Pt(10.5), "single", 0, 0, 0, False)
BODY_SPEC = StyleSpec(FONT_SONG, SIZE_BODY, "one_point_five", 0, 0, 2, None, WD_ALIGN_PARAGRAPH.JUSTIFY)
TABLE_SPEC = StyleSpec(FONT_SONG, SIZE_TABLE, "one_point_five", 0, 0, 0, None)
CODE_SPEC = StyleSpec(FONT_MONO, SIZE_CODE, "one_point_five", 6, 6, 0, None)
PAGE_SPEC = StyleSpec(FONT_SONG, Pt(10.5), "one_point_five", 0, 0, 0, None)
TOC_TITLE_SPEC = StyleSpec(FONT_HEI, SIZE_H1, "single", 10, 10, 0, True, font_color=BLACK_FONT_COLOR)
TOC_ENTRY_SPEC = StyleSpec(FONT_SONG, SIZE_BODY, "one_point_five", 0, 0, 0, None)

TOC_TOP_LEVEL_ENTRY_RE = re.compile(r'^\s*(\d+)\.\s*$')

HEADING_RE = re.compile(r'^(#{1,9})\s+(.*)$')
BULLET_RE = re.compile(r'^(\s*)[-*+]\s+(.*)$')
ORDERED_RE = re.compile(r'^\s*\d+[.)]\s+(.*)$')
LEADING_HEADING_NUMBER_RE = re.compile(r'^\s*((?:\d+\.)*\d+)\s+')
TXT_NUMBERED_HEADING_RE = re.compile(r'^\s*((?:\d+\.)*\d+)\s+(.+?)\s*$')
TXT_CHINESE_HEADING_RE = re.compile(r'^\s*([一二三四五六七八九十百千]+)、\s*(.+?)\s*$')
MARKDOWN_NUMBERED_HEADING_RE = re.compile(r'^\s*((?:\d+\.)*\d+)(?:[、.)]|\s)+(.+?)\s*$')
MARKDOWN_CHINESE_HEADING_RE = re.compile(r'^\s*([一二三四五六七八九十百千]+)、\s*(.+?)\s*$')
COVER_LABEL_RE = re.compile(r'^\s*(?:\[(封面标题|封面公司/日期|公司名称|日期)\]|(封面标题|封面公司/日期|公司名称|日期))\s*[:：]\s*(.*)$')
COVER_CORNER_LABEL_RE = re.compile(r'^\s*(编\s*号|版\s*本\s*号|受控状态|密\s*级)\s*[:：].*$')
DATE_RE = re.compile(r'\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b|\b\d{4}年\d{1,2}月\d{1,2}日\b')


class AnalysisResult:
    def __init__(self, cover, body_lines, toc_exists: bool, toc_entries: list[tuple[int, str]]):
        self.cover = cover
        self.body_lines = body_lines
        self.toc_exists = toc_exists
        self.toc_entries = toc_entries


class TxtBlock:
    def __init__(self, kind: str, lines: list[str], level: int | None = None, title: str | None = None):
        self.kind = kind
        self.lines = lines
        self.level = level
        self.title = title


class VerificationContext:
    def __init__(
        self,
        expected_cover: dict[str, list[str] | str | None] | None,
        expect_cover: bool,
        expect_toc: bool,
        expected_headings: list[tuple[int, str]],
        source_kind: str,
        toc_refresh_state: str = 'not_refreshed',
        structure_snapshot: dict[str, int] | None = None,
    ):
        self.expected_cover = expected_cover
        self.expect_cover = expect_cover
        self.expect_toc = expect_toc
        self.expected_headings = expected_headings
        self.source_kind = source_kind
        self.toc_refresh_state = toc_refresh_state
        self.structure_snapshot = structure_snapshot or {}


def clear_document_body(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)


def clear_header_footer_content(container):
    element = container._element
    for child in list(element):
        if child.tag != qn('w:pPr'):
            element.remove(child)


def apply_section_page_setup(section):
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.top_margin = PAGE_MARGIN_TOP
    section.bottom_margin = PAGE_MARGIN_BOTTOM
    section.left_margin = PAGE_MARGIN_LEFT
    section.right_margin = PAGE_MARGIN_RIGHT


def set_section_page_number_start(section, start: int | None):
    sect_pr = section._sectPr
    for child in list(sect_pr):
        if child.tag == qn('w:pgNumType'):
            sect_pr.remove(child)
    if start is None:
        return
    pg_num_type = OxmlElement('w:pgNumType')
    pg_num_type.set(qn('w:start'), str(start))
    sect_pr.append(pg_num_type)


def configure_section_footer(section, show_page_number: bool, page_number_start: int | None = None):
    section.different_first_page_header_footer = False
    set_section_page_number_start(section, page_number_start)

    for footer in (section.footer, section.first_page_footer, section.even_page_footer):
        footer.is_linked_to_previous = False
        clear_header_footer_content(footer)

    if not show_page_number:
        return

    paragraph = section.footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    apply_line_spacing(paragraph, PAGE_SPEC.line_spacing)
    add_page_number(paragraph)


def ensure_section_type_next_page(sect_pr) -> None:
    type_node = sect_pr.get_or_add_type()
    type_node.set(qn('w:val'), 'nextPage')


def start_body_section(doc: Document):
    section = doc.add_section()
    apply_section_page_setup(section)
    configure_section_footer(section, show_page_number=True, page_number_start=1)
    return section


def get_heading_level(source_level: int, heading_base_level: int) -> int:
    return min(max(source_level - heading_base_level + 1, 1), MAX_HEADING_LEVEL)


def get_heading_style_name(level: int) -> str:
    return f'Heading {min(level, 9)}'


def get_heading_style_id(level: int) -> str:
    return get_heading_style_name(level)


def resolve_heading_style_id(doc: Document, level: int) -> str:
    style = doc.styles[get_heading_style_name(level)]
    return style.style_id


def set_rpr_color(r_pr, color: str | None) -> None:
    if color is None:
        return
    color_node = r_pr.find(qn('w:color'))
    if color_node is None:
        color_node = OxmlElement('w:color')
        inserted = False
        after_tags = {
            qn('w:spacing'),
            qn('w:w'),
            qn('w:kern'),
            qn('w:position'),
            qn('w:sz'),
            qn('w:szCs'),
            qn('w:lang'),
        }
        for idx, child in enumerate(list(r_pr)):
            if child.tag in after_tags:
                r_pr.insert(idx, color_node)
                inserted = True
                break
        if not inserted:
            r_pr.append(color_node)
    color_node.set(qn('w:val'), color)


def get_or_create_paragraph_style(doc: Document, style_name: str, style_id: str, spec: StyleSpec):
    if style_name in doc.styles:
        style = doc.styles[style_name]
    elif style_id in doc.styles:
        style = doc.styles[style_id]
    else:
        raise KeyError(f'Style not found: {style_name} / {style_id}')

    style.font.name = spec.font_name
    style.font.size = spec.font_size
    style.font.bold = bool(spec.bold)
    apply_style_paragraph_format(style, spec)

    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), spec.font_name)
    r_fonts.set(qn('w:hAnsi'), spec.font_name)
    r_fonts.set(qn('w:eastAsia'), spec.font_name)
    set_rpr_color(r_pr, spec.font_color)

    return style


def set_run_font(run, font_name: str, font_size, font_color: str | None = None):
    run.font.name = font_name
    run.font.size = font_size
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)
    r_fonts.set(qn('w:eastAsia'), font_name)
    set_rpr_color(r_pr, font_color)


def insert_numpr_after_pstyle(p_pr, num_pr):
    inserted = False
    for idx, child in enumerate(list(p_pr)):
        if child.tag == qn('w:pStyle'):
            p_pr.insert(idx + 1, num_pr)
            inserted = True
            break
    if not inserted:
        p_pr.insert(0, num_pr)


def clear_paragraph_indentation(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn('w:ind'):
            p_pr.remove(child)


def clear_paragraph_spacing(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn('w:spacing'):
            p_pr.remove(child)


def clear_paragraph_alignment(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn('w:jc'):
            p_pr.remove(child)


def clear_paragraph_numbering(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn('w:numPr'):
            p_pr.remove(child)


def clear_style_numbering(style):
    p_pr = style._element.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn('w:numPr'):
            p_pr.remove(child)


def clear_paragraph_tabs(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn('w:tabs'):
            p_pr.remove(child)


def clear_paragraph_formatting(paragraph):
    clear_paragraph_indentation(paragraph)
    clear_paragraph_spacing(paragraph)
    clear_paragraph_alignment(paragraph)
    clear_paragraph_numbering(paragraph)
    clear_paragraph_tabs(paragraph)


def set_first_line_indent_chars(paragraph, chars: int):
    clear_paragraph_indentation(paragraph)
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.get_or_add_ind()
    ind.set(qn('w:left'), '0')
    ind.set(qn('w:leftChars'), '0')
    ind.set(qn('w:right'), '0')
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), str(chars * 100))


def set_no_first_line_indent(paragraph):
    clear_paragraph_indentation(paragraph)
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.get_or_add_ind()
    ind.set(qn('w:left'), '0')
    ind.set(qn('w:leftChars'), '0')
    ind.set(qn('w:right'), '0')
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), '0')


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn('w:shd'):
            p_pr.remove(child)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    p_pr.append(shd)


def apply_line_spacing(paragraph, mode: str):
    fmt = paragraph.paragraph_format
    if mode == 'single':
        fmt.line_spacing = 1.0
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif mode == 'one_point_five':
        fmt.line_spacing = 1.5
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:
        raise ValueError(f'Unsupported line spacing mode: {mode}')


def apply_style_paragraph_format(style, spec: StyleSpec):
    fmt = style.paragraph_format
    fmt.space_before = Pt(spec.before_pt)
    fmt.space_after = Pt(spec.after_pt)
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)
    p_pr = style._element.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn('w:ind'):
            p_pr.remove(child)
    ind = p_pr.get_or_add_ind()
    ind.set(qn('w:left'), '0')
    ind.set(qn('w:leftChars'), '0')
    ind.set(qn('w:right'), '0')
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), '0')
    if spec.line_spacing == 'single':
        fmt.line_spacing = 1.0
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif spec.line_spacing == 'one_point_five':
        fmt.line_spacing = 1.5
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:
        raise ValueError(f'Unsupported line spacing mode: {spec.line_spacing}')


def apply_paragraph_style(paragraph, spec: StyleSpec):
    clear_paragraph_formatting(paragraph)
    paragraph.alignment = spec.alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(spec.before_pt)
    fmt.space_after = Pt(spec.after_pt)
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    apply_line_spacing(paragraph, spec.line_spacing)
    if spec.first_line_chars > 0:
        set_first_line_indent_chars(paragraph, spec.first_line_chars)
    else:
        set_no_first_line_indent(paragraph)

    for run in paragraph.runs:
        set_run_font(run, spec.font_name, spec.font_size, spec.font_color)
        if spec.bold is not None:
            run.bold = spec.bold


def apply_code_block_style(paragraph):
    apply_paragraph_style(paragraph, CODE_SPEC)
    paragraph.paragraph_format.left_indent = CODE_LEFT_INDENT
    paragraph.paragraph_format.right_indent = CODE_RIGHT_INDENT
    set_paragraph_shading(paragraph, CODE_SHADE)


def get_heading_spec(level: int) -> StyleSpec:
    if level == 1:
        return H1_SPEC
    if level == 2:
        return H2_SPEC
    if level == 3:
        return H3_SPEC
    if level == 4:
        return H4_SPEC
    if level == 5:
        return H5_SPEC
    if level == 6:
        return H6_SPEC
    if level == 7:
        return H7_SPEC
    if level == 8:
        return H8_SPEC
    return H9_SPEC


def apply_bullet_style(paragraph, level: int):
    apply_paragraph_style(paragraph, BODY_SPEC)
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn('w:numPr'):
            p_pr.remove(child)

    num_pr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(max(min(level, 2), 0)))
    num_id = OxmlElement('w:numId')
    num_id.set(qn('w:val'), BULLET_NUMBERING_ID)
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)


def set_cell_margins(cell, top=TABLE_CELL_MARGIN_TOP, start=TABLE_CELL_MARGIN_LEFT, bottom=TABLE_CELL_MARGIN_BOTTOM, end=TABLE_CELL_MARGIN_RIGHT):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for tag, value in [('w:top', top), ('w:start', start), ('w:bottom', bottom), ('w:end', end)]:
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def add_field(paragraph, instruction: str, display_text: str | None = None, spec: StyleSpec = PAGE_SPEC):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instruction
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    if display_text is not None:
        fld_text = OxmlElement('w:t')
        fld_text.text = display_text
        run._r.append(fld_text)
    run._r.append(fld_end)
    set_run_font(run, spec.font_name, spec.font_size)


def add_page_number(paragraph):
    add_field(paragraph, ' PAGE ', '1', PAGE_SPEC)


def add_toc_field(paragraph, levels: str = '1-6'):
    add_field(paragraph, f' TOC \\o "{levels}" \\h \\z \\u ', '右键更新目录', TOC_ENTRY_SPEC)


def ensure_numbering_child_order(numbering, child, kind: str):
    children = list(numbering)
    if kind == 'abstract':
        for idx, existing in enumerate(children):
            if existing.tag in {qn('w:num'), qn('w:numIdMacAtCleanup')}:
                numbering.insert(idx, child)
                return
        numbering.append(child)
        return

    if kind == 'num':
        for idx, existing in enumerate(children):
            if existing.tag == qn('w:numIdMacAtCleanup'):
                numbering.insert(idx, child)
                return
        numbering.append(child)
        return

    raise ValueError(f'Unsupported numbering child kind: {kind}')


def ensure_title_numbering(doc: Document, numbering, abstract_id: str, numbering_id: str):
    title_existing = numbering.xpath(f'./w:abstractNum[@w:abstractNumId="{abstract_id}"]')
    if not title_existing:
        title_abstract = OxmlElement('w:abstractNum')
        title_abstract.set(qn('w:abstractNumId'), abstract_id)

        nsid = OxmlElement('w:nsid')
        nsid.set(qn('w:val'), '6C6F6E67')
        title_abstract.append(nsid)

        multi_level_type = OxmlElement('w:multiLevelType')
        multi_level_type.set(qn('w:val'), 'multilevel')
        title_abstract.append(multi_level_type)

        level_styles = {
            ilvl: resolve_heading_style_id(doc, ilvl + 1) for ilvl in range(MAX_HEADING_LEVEL)
        }
        level_patterns = {
            ilvl: '.'.join(f'%{idx}' for idx in range(1, ilvl + 2))
            for ilvl in range(MAX_HEADING_LEVEL)
        }

        for ilvl in range(MAX_HEADING_LEVEL):
            lvl = OxmlElement('w:lvl')
            lvl.set(qn('w:ilvl'), str(ilvl))

            start = OxmlElement('w:start')
            start.set(qn('w:val'), '1')
            lvl.append(start)

            num_fmt = OxmlElement('w:numFmt')
            num_fmt.set(qn('w:val'), 'decimal')
            lvl.append(num_fmt)

            p_style = OxmlElement('w:pStyle')
            p_style.set(qn('w:val'), level_styles[ilvl])
            lvl.append(p_style)

            suff = OxmlElement('w:suff')
            suff.set(qn('w:val'), 'space')
            lvl.append(suff)

            lvl_text = OxmlElement('w:lvlText')
            lvl_text.set(qn('w:val'), level_patterns[ilvl])
            lvl.append(lvl_text)

            lvl_jc = OxmlElement('w:lvlJc')
            lvl_jc.set(qn('w:val'), 'left')
            lvl.append(lvl_jc)

            r_pr = OxmlElement('w:rPr')
            r_fonts = OxmlElement('w:rFonts')
            spec = get_heading_spec(ilvl + 1)
            r_fonts.set(qn('w:ascii'), spec.font_name)
            r_fonts.set(qn('w:hAnsi'), spec.font_name)
            r_fonts.set(qn('w:eastAsia'), spec.font_name)
            r_pr.append(r_fonts)
            set_rpr_color(r_pr, spec.font_color)
            if spec.bold:
                r_pr.append(OxmlElement('w:b'))
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(spec.font_size.pt * 2)))
            r_pr.append(sz)
            sz_cs = OxmlElement('w:szCs')
            sz_cs.set(qn('w:val'), str(int(spec.font_size.pt * 2)))
            r_pr.append(sz_cs)
            lvl.append(r_pr)

            title_abstract.append(lvl)

        ensure_numbering_child_order(numbering, title_abstract, 'abstract')

    title_num_existing = numbering.xpath(f'./w:num[@w:numId="{numbering_id}"]')
    if not title_num_existing:
        title_num = OxmlElement('w:num')
        title_num.set(qn('w:numId'), numbering_id)
        title_ref = OxmlElement('w:abstractNumId')
        title_ref.set(qn('w:val'), abstract_id)
        title_num.append(title_ref)
        ensure_numbering_child_order(numbering, title_num, 'num')


def ensure_numbering(doc: Document):
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    ensure_title_numbering(doc, numbering, TITLE_NUMBERING_ABSTRACT_ID, TITLE_NUMBERING_ID)

    bullet_existing = numbering.xpath(f'./w:abstractNum[@w:abstractNumId="{BULLET_NUMBERING_ABSTRACT_ID}"]')
    if bullet_existing:
        return

    bullet_abstract = OxmlElement('w:abstractNum')
    bullet_abstract.set(qn('w:abstractNumId'), BULLET_NUMBERING_ABSTRACT_ID)
    bullet_chars = ['●', '■', '◆']
    for ilvl, char in enumerate(bullet_chars):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'bullet')
        lvl.append(num_fmt)

        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), char)
        lvl.append(lvl_text)

        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)

        p_pr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '0')
        ind.set(qn('w:hanging'), '0')
        p_pr.append(ind)
        lvl.append(p_pr)

        r_pr = OxmlElement('w:rPr')
        r_fonts = OxmlElement('w:rFonts')
        r_fonts.set(qn('w:ascii'), FONT_HEI)
        r_fonts.set(qn('w:hAnsi'), FONT_HEI)
        r_fonts.set(qn('w:eastAsia'), FONT_HEI)
        r_pr.append(r_fonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(SIZE_BODY.pt * 2)))
        r_pr.append(sz)
        sz_cs = OxmlElement('w:szCs')
        sz_cs.set(qn('w:val'), str(int(SIZE_BODY.pt * 2)))
        r_pr.append(sz_cs)
        lvl.append(r_pr)
        bullet_abstract.append(lvl)

    ensure_numbering_child_order(numbering, bullet_abstract, 'abstract')

    bullet_num = OxmlElement('w:num')
    bullet_num.set(qn('w:numId'), BULLET_NUMBERING_ID)
    bullet_ref = OxmlElement('w:abstractNumId')
    bullet_ref.set(qn('w:val'), BULLET_NUMBERING_ABSTRACT_ID)
    bullet_num.append(bullet_ref)
    ensure_numbering_child_order(numbering, bullet_num, 'num')


def get_title_numbering_id() -> str:
    return TITLE_NUMBERING_ID


def apply_heading_numbering(paragraph, level: int):
    clear_paragraph_numbering(paragraph)

    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(max(level - 1, 0)))
    num_id = OxmlElement('w:numId')
    num_id.set(qn('w:val'), TITLE_NUMBERING_ID)
    num_pr.append(ilvl)
    num_pr.append(num_id)
    insert_numpr_after_pstyle(p_pr, num_pr)


def normalize_heading_text(text: str) -> str:
    content = text.strip()
    chinese = MARKDOWN_CHINESE_HEADING_RE.match(content)
    if chinese:
        return chinese.group(2).strip()
    numbered = MARKDOWN_NUMBERED_HEADING_RE.match(content)
    if numbered:
        return numbered.group(2).strip()
    return strip_heading_number_prefix(content)


def get_heading_levels_present(entries: list[tuple[int, str]]) -> list[int]:
    return sorted({level for level, title in entries if title})


def get_toc_field_levels(entries: list[tuple[int, str]]) -> str:
    levels = get_heading_levels_present(entries)
    if not levels:
        return '1-3'
    return f'1-{max(levels)}'


def get_expected_heading_number_pattern(level: int) -> str:
    return '.'.join(f'%{idx}' for idx in range(1, level + 1))


def extract_paragraph_numpr(paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None
    for child in list(p_pr):
        if child.tag == qn('w:numPr'):
            return child
    return None


def get_numpr_values(num_pr) -> tuple[int | None, str | None]:
    ilvl_value = None
    num_id_value = None
    for child in list(num_pr):
        if child.tag == qn('w:ilvl'):
            raw = child.get(qn('w:val'))
            if raw is not None:
                ilvl_value = int(raw)
        elif child.tag == qn('w:numId'):
            num_id_value = child.get(qn('w:val'))
    return ilvl_value, num_id_value


def get_style_numpr(style):
    p_pr = style._element.pPr
    if p_pr is None:
        return None
    for child in list(p_pr):
        if child.tag == qn('w:numPr'):
            return child
    return None


def get_numbering_nodes(numbering):
    abstract_by_id: dict[str, object] = {}
    num_by_id: dict[str, object] = {}
    for child in list(numbering):
        if child.tag == qn('w:abstractNum'):
            abstract_id = child.get(qn('w:abstractNumId'))
            if abstract_id is not None:
                abstract_by_id[abstract_id] = child
        elif child.tag == qn('w:num'):
            num_id = child.get(qn('w:numId'))
            if num_id is not None:
                num_by_id[num_id] = child
    return abstract_by_id, num_by_id


def get_num_abstract_id(num_node) -> str | None:
    for child in list(num_node):
        if child.tag == qn('w:abstractNumId'):
            return child.get(qn('w:val'))
    return None


def detach_inactive_heading_style_links(doc: Document, active_numbering_id: str) -> None:
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return
    numbering = numbering_part.numbering_definitions._numbering
    abstract_by_id, num_by_id = get_numbering_nodes(numbering)
    num_node = num_by_id.get(active_numbering_id)
    if num_node is None:
        return
    active_abstract_id = get_num_abstract_id(num_node)
    if active_abstract_id is None:
        return
    heading_style_ids = {resolve_heading_style_id(doc, level) for level in range(1, MAX_HEADING_LEVEL + 1)}
    for abstract_id, abstract_node in abstract_by_id.items():
        if abstract_id == active_abstract_id:
            continue
        for level_node in list(abstract_node):
            if level_node.tag != qn('w:lvl'):
                continue
            for child in list(level_node):
                if child.tag == qn('w:pStyle') and child.get(qn('w:val')) in heading_style_ids:
                    level_node.remove(child)


def get_abstract_level_node(abstract_node, ilvl: int):
    for child in list(abstract_node):
        if child.tag == qn('w:lvl') and child.get(qn('w:ilvl')) == str(ilvl):
            return child
    return None


def get_numbering_level_node(doc: Document, num_id: str, ilvl: int):
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return None
    numbering = numbering_part.numbering_definitions._numbering
    abstract_by_id, num_by_id = get_numbering_nodes(numbering)
    num_node = num_by_id.get(num_id)
    if num_node is None:
        return None
    abstract_id = get_num_abstract_id(num_node)
    if abstract_id is None:
        return None
    abstract_node = abstract_by_id.get(abstract_id)
    if abstract_node is None:
        return None
    return get_abstract_level_node(abstract_node, ilvl)


def get_level_text_value(level_node) -> str | None:
    for child in list(level_node):
        if child.tag == qn('w:lvlText'):
            return child.get(qn('w:val'))
    return None


def get_outline_level(style) -> int | None:
    p_pr = style._element.pPr
    if p_pr is None:
        return None
    for child in list(p_pr):
        if child.tag == qn('w:outlineLvl'):
            raw = child.get(qn('w:val'))
            if raw is not None:
                return int(raw)
    return None


def collect_heading_paragraphs(doc: Document) -> dict[int, list[object]]:
    result: dict[int, list[object]] = {}
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ''
        if not style_name.startswith('Heading '):
            continue
        try:
            level = int(style_name.split(' ')[1])
        except (IndexError, ValueError):
            continue
        result.setdefault(level, []).append(paragraph)
    return result


def get_paragraph_style_id(paragraph) -> str | None:
    p_pr = paragraph._p.pPr
    if p_pr is not None:
        for child in list(p_pr):
            if child.tag == qn('w:pStyle'):
                return child.get(qn('w:val'))
    style = paragraph.style
    if style is not None and style.style_id is not None:
        return style.style_id
    return None


def iter_paragraph_text_nodes(paragraph):
    for child in paragraph._p.iter():
        if child.tag == qn('w:t'):
            yield child


def chinese_counting_text(value: int, force_one_ten: bool = False) -> str:
    if value <= 0:
        raise ValueError('Chinese counting value must be positive')
    digits = '零一二三四五六七八九'
    if value < 10:
        return digits[value]
    if value < 20:
        suffix = digits[value % 10] if value % 10 else ''
        return f'一十{suffix}' if force_one_ten else f'十{suffix}'
    if value < 100:
        tens, ones = divmod(value, 10)
        suffix = digits[ones] if ones else ''
        return f'{digits[tens]}十{suffix}'
    if value < 1000:
        hundreds, remainder = divmod(value, 100)
        prefix = f'{digits[hundreds]}百'
        if remainder == 0:
            return prefix
        if remainder < 10:
            return f'{prefix}零{chinese_counting_text(remainder, force_one_ten=True)}'
        return f'{prefix}{chinese_counting_text(remainder, force_one_ten=True)}'
    thousands, remainder = divmod(value, 1000)
    prefix = f'{digits[thousands]}千'
    if remainder == 0:
        return prefix
    if remainder < 100:
        return f'{prefix}零{chinese_counting_text(remainder, force_one_ten=True)}'
    return f'{prefix}{chinese_counting_text(remainder, force_one_ten=True)}'


def refresh_existing_toc_cache(doc: Document) -> bool:
    return False


def clear_toc_paragraph_indentation(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for child in list(p_pr):
        if child.tag == qn('w:ind'):
            p_pr.remove(child)


def normalize_existing_toc_paragraphs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        style_id = get_paragraph_style_id(paragraph)
        if style_id is None or not style_id.upper().startswith('TOC'):
            continue
        clear_toc_paragraph_indentation(paragraph)


def get_paragraph_text(paragraph) -> str:
    return ''.join((text_node.text or '') for text_node in iter_paragraph_text_nodes(paragraph)).strip()


def get_paragraph_section_index(doc: Document, paragraph) -> int:
    paragraph_element = paragraph._p
    section_index = 0
    for child in doc._element.body.iterchildren():
        if child == paragraph_element:
            return section_index
        if child.tag == qn('w:sectPr'):
            section_index += 1
    return section_index


def get_non_empty_paragraphs(paragraphs: list[object]) -> list[object]:
    return [paragraph for paragraph in paragraphs if get_paragraph_text(paragraph)]


def get_run_font_names(paragraph) -> set[str]:
    return {run.font.name for run in paragraph.runs if run.text and run.font.name is not None}


def get_run_font_sizes(paragraph) -> set[float]:
    sizes: set[float] = set()
    for run in paragraph.runs:
        if not run.text or run.font.size is None:
            continue
        sizes.add(run.font.size.pt)
    return sizes


def get_run_bold_values(paragraph) -> set[bool]:
    return {run.bold for run in paragraph.runs if run.text and run.bold is not None}


def is_table_paragraph(doc: Document, paragraph) -> bool:
    parent = paragraph._p.getparent()
    while parent is not None:
        if parent.tag == qn('w:tc'):
            return True
        parent = parent.getparent()
    return False


def is_code_paragraph(paragraph) -> bool:
    font_names = get_run_font_names(paragraph)
    if FONT_MONO in font_names:
        return True
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    return any(child.tag == qn('w:shd') for child in list(p_pr))


def find_first_toc_heading_index(doc: Document) -> int | None:
    for idx, paragraph in enumerate(doc.paragraphs):
        if is_explicit_toc_heading(get_paragraph_text(paragraph)):
            return idx
    return None


def collect_cover_paragraphs(doc: Document) -> list[object]:
    toc_index = find_first_toc_heading_index(doc)
    cover_paragraphs: list[object] = []
    title_started = False
    title_ended = False
    for idx, paragraph in enumerate(doc.paragraphs):
        if toc_index is not None and idx >= toc_index:
            break
        if paragraph.style is not None and paragraph.style.name.startswith('Heading '):
            break
        text = get_paragraph_text(paragraph)
        style_name = paragraph.style.name if paragraph.style is not None else ''
        is_corner = bool(text and COVER_CORNER_LABEL_RE.match(text))
        is_title = bool(text and (style_name.startswith('文章标题') or (paragraph.alignment is not None and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER)))
        if is_title and not title_ended:
            title_started = True
            cover_paragraphs.append(paragraph)
            continue
        if is_corner and not title_started:
            cover_paragraphs.append(paragraph)
            continue
        if title_started and not title_ended and not text:
            cover_paragraphs.append(paragraph)
            continue
        if title_started and not is_title:
            title_ended = True
            break
        if not title_started:
            if text:
                break
            cover_paragraphs.append(paragraph)
    return cover_paragraphs


def get_non_empty_cover_paragraphs(doc: Document) -> list[object]:
    return get_non_empty_paragraphs(collect_cover_paragraphs(doc))


def find_toc_region(doc: Document) -> tuple[object | None, object | None]:
    toc_heading = None
    toc_field = None
    for paragraph in doc.paragraphs:
        text = get_paragraph_text(paragraph)
        if toc_heading is None and is_explicit_toc_heading(text):
            toc_heading = paragraph
            continue
        if toc_heading is not None and toc_field is None and any(child.tag == qn('w:fldChar') for run in paragraph.runs for child in run._r):
            toc_field = paragraph
            break
    return toc_heading, toc_field


def get_first_body_content_child(doc: Document):
    for child in doc._element.body.iterchildren():
        if child.tag != qn('w:sectPr'):
            return child
    return None


def remove_body_child_range(doc: Document, start, stop=None) -> None:
    body = doc._element.body
    removing = start is None
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        if not removing:
            if child != start:
                continue
            removing = True
        if stop is not None and child == stop:
            break
        body.remove(child)


def trim_leading_empty_body_paragraphs(doc: Document) -> None:
    body = doc._element.body
    while True:
        first = get_first_body_content_child(doc)
        if first is None or first.tag == qn('w:tbl'):
            return
        if first.tag != qn('w:p'):
            return
        paragraph = DocxParagraph(first, doc._body)
        if get_paragraph_text(paragraph):
            return
        body.remove(first)


def remove_existing_toc(doc: Document) -> None:
    toc_heading, toc_field = find_toc_region(doc)
    if toc_heading is None and toc_field is None:
        return
    start = toc_heading._p if toc_heading is not None else toc_field._p
    boundary = find_body_start_boundary(doc)
    stop = None if boundary == start else boundary
    remove_body_child_range(doc, start, stop)
    trim_leading_empty_body_paragraphs(doc)


def remove_existing_cover(doc: Document) -> None:
    cover_paragraphs = collect_cover_paragraphs(doc)
    if not cover_paragraphs:
        return
    for paragraph in cover_paragraphs:
        p_element = paragraph._p
        parent = p_element.getparent()
        if parent is not None:
            parent.remove(p_element)
    trim_leading_empty_body_paragraphs(doc)


def find_body_start_boundary(doc: Document) -> object | None:
    toc_heading, toc_field = find_toc_region(doc)
    if toc_heading is not None or toc_field is not None:
        anchor = toc_field._p if toc_field is not None else toc_heading._p
        passed_anchor = False
        for child in doc._element.body.iterchildren():
            if child.tag == qn('w:sectPr'):
                continue
            if not passed_anchor:
                if child == anchor:
                    passed_anchor = True
                continue
            if child.tag == qn('w:p'):
                paragraph = DocxParagraph(child, doc._body)
                if get_paragraph_text(paragraph):
                    return child
            elif child.tag == qn('w:tbl'):
                return child
        return anchor
    for child in doc._element.body.iterchildren():
        if child.tag == qn('w:p'):
            paragraph = DocxParagraph(child, doc._body)
            style_name = paragraph.style.name if paragraph.style is not None else ''
            if style_name.startswith('Heading '):
                return child
    for child in doc._element.body.iterchildren():
        if child.tag == qn('w:p'):
            paragraph = DocxParagraph(child, doc._body)
            if get_paragraph_text(paragraph):
                return child
        elif child.tag == qn('w:tbl'):
            return child
    return None


def iter_body_blocks(doc: Document):
    body = doc._element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            yield DocxParagraph(child, doc._body)
            continue
        if child.tag == qn('w:tbl'):
            yield DocxTable(child, doc._body)


def infer_expected_headings(doc: Document) -> list[tuple[int, str]]:
    headings: list[tuple[int, str]] = []
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ''
        if not style_name.startswith('Heading '):
            continue
        try:
            level = int(style_name.split(' ')[1])
        except (IndexError, ValueError):
            continue
        text = get_paragraph_text(paragraph)
        if not text:
            continue
        headings.append((level, text))
    return headings


def infer_cover(doc: Document) -> dict[str, list[str] | str | None] | None:
    cover_paragraphs = []
    for paragraph in collect_cover_paragraphs(doc):
        text = get_paragraph_text(paragraph)
        if not text:
            continue
        cover_paragraphs.append(paragraph)
        if len(cover_paragraphs) >= 8:
            break
    if not cover_paragraphs:
        return None

    title = None
    title_lines: list[str] = []
    meta: list[str] = []
    corner_meta: list[str] = []
    title_candidates: list[str] = []
    for paragraph in cover_paragraphs:
        text = get_paragraph_text(paragraph)
        if COVER_CORNER_LABEL_RE.match(text):
            corner_meta.append(text)
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ''
        if style_name.startswith('文章标题') or style_name.startswith('Title') or style_name.startswith('Heading'):
            title_candidates.append(text)
            continue
        if paragraph.alignment is not None and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            title_candidates.append(text)
            continue
        meta.append(text)

    if title_candidates:
        title_lines = title_candidates
        title = title_candidates[0]
    elif meta:
        title = meta[0]
        title_lines = [meta[0]]
        meta = meta[1:]

    return {'title': title, 'title_lines': title_lines, 'meta': meta, 'corner_meta': corner_meta}


def count_descendants_by_tag(element, tag: str) -> int:
    return sum(1 for child in element.iter() if child.tag == tag)


def count_page_breaks(element) -> int:
    count = 0
    for child in element.iter():
        if child.tag != qn('w:br'):
            continue
        if child.get(qn('w:type')) == 'page':
            count += 1
    return count


def snapshot_document_structure(doc: Document) -> dict[str, int]:
    body = doc._element.body
    return {
        'drawing': count_descendants_by_tag(body, qn('w:drawing')),
        'pict': count_descendants_by_tag(body, qn('w:pict')),
        'object': count_descendants_by_tag(body, qn('w:object')),
        'table': count_descendants_by_tag(body, qn('w:tbl')),
        'page_break': count_page_breaks(body),
        'section': count_descendants_by_tag(body, qn('w:sectPr')),
    }


def apply_existing_cover_title_font(paragraph) -> None:
    for run in paragraph.runs:
        if not run.text:
            continue
        set_run_font(run, COVER_TITLE_SPEC.font_name, COVER_TITLE_SPEC.font_size)


def normalize_existing_cover(doc: Document, cover: dict[str, list[str] | str | None] | None) -> None:
    if not cover:
        return
    title_lines = set(cover.get('title_lines') or ([] if not cover.get('title') else [str(cover['title'])]))
    meta = set(cover.get('meta') or [])
    corner_meta = set(cover.get('corner_meta') or [])
    for paragraph in collect_cover_paragraphs(doc):
        text = get_paragraph_text(paragraph)
        if not text:
            continue
        if text in title_lines:
            apply_existing_cover_title_font(paragraph)
            continue
        if text in corner_meta:
            apply_paragraph_style(paragraph, COVER_CORNER_SPEC)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if text in meta:
            apply_paragraph_style(paragraph, COVER_META_SPEC)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def ensure_body_section_for_existing_docx(doc: Document) -> None:
    boundary = find_body_start_boundary(doc)
    if boundary is None:
        return

    body = doc._element.body
    body_children = [child for child in body.iterchildren() if child.tag != qn('w:sectPr')]
    try:
        boundary_index = body_children.index(boundary)
    except ValueError:
        return
    if boundary_index == 0:
        return

    previous = body_children[boundary_index - 1]
    if previous.tag == qn('w:p'):
        previous_p_pr = previous.get_or_add_pPr()
        for child in list(previous_p_pr):
            if child.tag == qn('w:sectPr'):
                return

    boundary_sect_pr = None
    if boundary.tag == qn('w:p'):
        boundary_p_pr = boundary.pPr
        if boundary_p_pr is not None:
            for child in list(boundary_p_pr):
                if child.tag == qn('w:sectPr'):
                    boundary_sect_pr = child
                    boundary_p_pr.remove(child)
                    break

    sect_pr = boundary_sect_pr if boundary_sect_pr is not None else OxmlElement('w:sectPr')
    ensure_section_type_next_page(sect_pr)

    if previous.tag == qn('w:p'):
        previous.get_or_add_pPr().append(sect_pr)
        return

    section_break_paragraph = OxmlElement('w:p')
    p_pr = OxmlElement('w:pPr')
    p_pr.append(sect_pr)
    section_break_paragraph.append(p_pr)
    body.insert(list(body).index(boundary), section_break_paragraph)


def normalize_existing_heading_paragraph(paragraph) -> None:
    style_name = paragraph.style.name if paragraph.style is not None else ''
    if not style_name.startswith('Heading '):
        return
    try:
        level = int(style_name.split(' ')[1])
    except (IndexError, ValueError):
        return
    paragraph.style = get_heading_style_id(level)
    spec = get_heading_spec(level)
    normalized = normalize_heading_text(get_paragraph_text(paragraph))
    if normalized and normalized != get_paragraph_text(paragraph):
        if paragraph.runs:
            paragraph.runs[0].text = normalized
            for run in paragraph.runs[1:]:
                run.text = ''
        else:
            paragraph.add_run(normalized)
    apply_paragraph_style(paragraph, spec)
    clear_paragraph_numbering(paragraph)
    apply_heading_numbering(paragraph, level)


def normalize_existing_body_paragraph(paragraph) -> None:
    if is_code_paragraph(paragraph):
        return
    apply_paragraph_style(paragraph, BODY_SPEC)


def normalize_existing_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if not get_paragraph_text(paragraph):
                    continue
                if is_code_paragraph(paragraph):
                    continue
                apply_paragraph_style(paragraph, TABLE_SPEC)
                for run in paragraph.runs:
                    if not run.text:
                        continue
                    set_run_font(run, TABLE_SPEC.font_name, TABLE_SPEC.font_size)


def normalize_existing_docx_body(doc: Document) -> list[tuple[int, str]]:
    boundary = find_body_start_boundary(doc)
    started = boundary is None
    expected_headings: list[tuple[int, str]] = []
    for block in iter_body_blocks(doc):
        element = block._element if hasattr(block, '_element') else None
        if not started:
            if element == boundary:
                started = True
            else:
                continue
        if isinstance(block, DocxParagraph):
            text = get_paragraph_text(block)
            if not text:
                continue
            if is_explicit_toc_heading(text) or text == '右键更新目录':
                continue
            style_name = block.style.name if block.style is not None else ''
            if style_name.startswith('Heading '):
                normalize_existing_heading_paragraph(block)
                try:
                    level = int(style_name.split(' ')[1])
                except (IndexError, ValueError):
                    continue
                normalized = get_paragraph_text(block)
                if normalized:
                    expected_headings.append((level, normalized))
                continue
            normalize_existing_body_paragraph(block)
            continue
        if isinstance(block, DocxTable):
            normalize_existing_table(block)
    return expected_headings


def refresh_existing_docx(
    input_path: Path,
    output_path: Path,
    force_cover: bool | None = None,
    force_toc: bool | None = None,
) -> None:
    doc = Document(str(input_path))
    if force_toc is False:
        remove_existing_toc(doc)
    if force_cover is False:
        remove_existing_cover(doc)
    configure_document_styles(doc)
    inferred_cover = None if force_cover is False else infer_cover(doc)
    has_cover_section = bool(inferred_cover)
    structure_snapshot = snapshot_document_structure(doc)
    normalize_existing_cover(doc, inferred_cover)
    if has_cover_section:
        ensure_body_section_for_existing_docx(doc)
    expected_headings = normalize_existing_docx_body(doc)
    for table in collect_tables(doc):
        normalize_existing_table(table)
    toc_cache_updated = False if force_toc is False else refresh_existing_toc_cache(doc)
    if force_toc is not False:
        normalize_existing_toc_paragraphs(doc)

    if has_cover_section:
        configure_section_footer(doc.sections[0], show_page_number=False, page_number_start=None)
        if len(doc.sections) < 2:
            configure_section_footer(doc.sections[0], show_page_number=False, page_number_start=None)
        else:
            configure_section_footer(doc.sections[-1], show_page_number=True, page_number_start=1)
    else:
        configure_section_footer(doc.sections[0], show_page_number=True, page_number_start=1)

    toc_heading, toc_field = find_toc_region(doc)
    context = build_verification_context(
        expected_cover=inferred_cover,
        expect_cover=has_cover_section,
        expect_toc=bool(toc_heading or toc_field),
        expected_headings=expected_headings,
        source_kind='refreshed',
        toc_refresh_state='refreshed' if toc_cache_updated else 'not_refreshed',
        structure_snapshot=structure_snapshot,
    )
    verify_document_workflow(doc, context)
    doc.save(output_path)


def collect_tables(doc: Document) -> list[object]:
    return list(doc.tables)


def collect_body_paragraphs(doc: Document, context: VerificationContext | None = None) -> list[object]:
    boundary = find_body_start_boundary(doc)
    result: list[object] = []
    cover_texts = set()
    if context is not None and context.expected_cover is not None:
        title = context.expected_cover.get('title')
        if title:
            cover_texts.add(str(title))
        for item in context.expected_cover.get('title_lines') or []:
            cover_texts.add(str(item))
        for item in context.expected_cover.get('meta') or []:
            cover_texts.add(str(item))
        for item in context.expected_cover.get('corner_meta') or []:
            cover_texts.add(str(item))

    started = boundary is None
    for paragraph in doc.paragraphs:
        if not started:
            if paragraph._p == boundary:
                started = True
            else:
                continue
        text = get_paragraph_text(paragraph)
        if not text:
            continue
        if is_table_paragraph(doc, paragraph):
            continue
        if is_code_paragraph(paragraph):
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ''
        if style_name.startswith('Heading '):
            continue
        if is_explicit_toc_heading(text):
            continue
        if text == '右键更新目录':
            continue
        if text in cover_texts:
            continue
        if COVER_CORNER_LABEL_RE.match(text):
            continue
        result.append(paragraph)
    return result


def get_paragraph_spacing_points(paragraph) -> tuple[float | None, float | None]:
    before = paragraph.paragraph_format.space_before
    after = paragraph.paragraph_format.space_after
    before_pt = before.pt if before is not None else None
    after_pt = after.pt if after is not None else None
    return before_pt, after_pt


def get_style_spacing_points(style) -> tuple[float | None, float | None]:
    before = style.paragraph_format.space_before
    after = style.paragraph_format.space_after
    before_pt = before.pt if before is not None else None
    after_pt = after.pt if after is not None else None
    return before_pt, after_pt


def format_failures(prefix: str, failures: list[str]) -> list[str]:
    return [f'{prefix}: {failure}' for failure in failures]


def build_verification_context(
    expected_cover: dict[str, list[str] | str | None] | None,
    expect_cover: bool,
    expect_toc: bool,
    expected_headings: list[tuple[int, str]],
    source_kind: str,
    toc_refresh_state: str = 'not_refreshed',
    structure_snapshot: dict[str, int] | None = None,
) -> VerificationContext:
    return VerificationContext(
        expected_cover=expected_cover,
        expect_cover=expect_cover,
        expect_toc=expect_toc,
        expected_headings=expected_headings,
        source_kind=source_kind,
        toc_refresh_state=toc_refresh_state,
        structure_snapshot=structure_snapshot,
    )


def verify_heading_style_definitions(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    present_levels = get_heading_levels_present(context.expected_headings)
    for level in present_levels:
        spec = get_heading_spec(level)
        style_name = get_heading_style_name(level)
        if style_name not in doc.styles:
            failures.append(f'Heading {level} style "{style_name}" missing from document styles')
            continue
        style = doc.styles[style_name]
        if style.font.name != spec.font_name or style.font.size != spec.font_size or bool(style.font.bold) != bool(spec.bold):
            failures.append(f'Heading {level} style font mismatch')
        if get_outline_level(style) != min(level, 9) - 1:
            failures.append(f'Heading {level} outline level mismatch')
        before_pt, after_pt = get_style_spacing_points(style)
        if before_pt != spec.before_pt or after_pt != spec.after_pt:
            failures.append(f'Heading {level} style spacing mismatch')
        style_num_pr = get_style_numpr(style)
        if style_num_pr is None:
            failures.append(f'Heading {level} style missing numbering')
            continue
        ilvl_value, _ = get_numpr_values(style_num_pr)
        if ilvl_value != level - 1:
            failures.append(f'Heading {level} style numbering level mismatch')
    return failures


def verify_heading_paragraph_instances(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    heading_paragraphs = collect_heading_paragraphs(doc)
    expected_titles_by_level: dict[int, list[str]] = {}
    for level, title in context.expected_headings:
        expected_titles_by_level.setdefault(level, []).append(title)

    for level, expected_titles in expected_titles_by_level.items():
        spec = get_heading_spec(level)
        paragraphs = heading_paragraphs.get(level, [])
        if len(paragraphs) != len(expected_titles):
            failures.append(f'Heading {level} count mismatch: body={len(paragraphs)} expected={len(expected_titles)}')
        for idx, expected_title in enumerate(expected_titles):
            if idx >= len(paragraphs):
                break
            paragraph = paragraphs[idx]
            visible = get_paragraph_text(paragraph)
            if visible != expected_title:
                failures.append(f'Heading {level} text mismatch: body="{visible}" expected="{expected_title}"')
            if MARKDOWN_NUMBERED_HEADING_RE.match(visible) or MARKDOWN_CHINESE_HEADING_RE.match(visible):
                failures.append(f'Heading {level} still contains manual numbering prefix: {visible}')
            if any(run.font.name not in {None, spec.font_name} for run in paragraph.runs if run.text):
                failures.append(f'Heading {level} run font mismatch: {visible}')
            if any(run.font.size not in {None, spec.font_size} for run in paragraph.runs if run.text):
                failures.append(f'Heading {level} run size mismatch: {visible}')
            if spec.bold is not None and any(run.bold not in {None, spec.bold} for run in paragraph.runs if run.text):
                failures.append(f'Heading {level} run bold mismatch: {visible}')
    return failures


def verify_heading_numbering_definitions(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    for level, expected_title in context.expected_headings:
        paragraphs = collect_heading_paragraphs(doc).get(level, [])
        matches = [paragraph for paragraph in paragraphs if get_paragraph_text(paragraph) == expected_title]
        if not matches:
            continue
        paragraph = matches[0]
        num_pr = extract_paragraph_numpr(paragraph)
        if num_pr is None:
            failures.append(f'Heading {level} paragraph missing numbering: {expected_title}')
            continue
        ilvl_value, num_id_value = get_numpr_values(num_pr)
        if ilvl_value != level - 1:
            failures.append(f'Heading {level} paragraph numbering level mismatch: {expected_title}')
            continue
        if num_id_value is None:
            failures.append(f'Heading {level} paragraph numbering id missing: {expected_title}')
            continue
        level_node = get_numbering_level_node(doc, num_id_value, ilvl_value)
        if level_node is None:
            failures.append(f'Heading {level} numbering definition missing: {expected_title}')
            continue
        expected_pattern = get_expected_heading_number_pattern(level)
        if get_level_text_value(level_node) != expected_pattern:
            failures.append(f'Heading {level} numbering pattern mismatch: expected {expected_pattern}')
    return failures


def verify_heading_numbering_chain(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    previous_level = None
    for level, title in context.expected_headings:
        if previous_level is not None and level > previous_level + 1:
            failures.append(f'Heading level jump detected before: {title}')
        previous_level = level
    return failures


def verify_cover_presence(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    cover_paragraphs = get_non_empty_cover_paragraphs(doc)
    has_cover = bool(cover_paragraphs)
    if context.expect_cover and not has_cover:
        failures.append('cover expected but no non-empty cover paragraphs found')
    if not context.expect_cover and has_cover:
        failures.append('cover not expected but cover-like paragraphs were retained')
    return failures


def verify_cover_title_block(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    if not context.expect_cover or context.expected_cover is None:
        return failures
    title_lines = list(context.expected_cover.get('title_lines') or ([] if not context.expected_cover.get('title') else [str(context.expected_cover.get('title'))]))
    if not title_lines:
        return failures
    cover_paragraphs = get_non_empty_cover_paragraphs(doc)
    by_text = {get_paragraph_text(paragraph): paragraph for paragraph in cover_paragraphs}
    for title in title_lines:
        paragraph = by_text.get(title)
        if paragraph is None:
            failures.append(f'cover title missing: {title}')
            continue
        if get_run_font_names(paragraph) not in [set(), {COVER_TITLE_SPEC.font_name}]:
            failures.append(f'cover title font mismatch: {title}')
        if get_run_font_sizes(paragraph) not in [set(), {COVER_TITLE_SPEC.font_size.pt}]:
            failures.append(f'cover title size mismatch: {title}')
        if context.source_kind != 'refreshed':
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                failures.append(f'cover title is not centered: {title}')
            if get_run_bold_values(paragraph) not in [set(), {True}]:
                failures.append(f'cover title bold mismatch: {title}')
            before_pt, after_pt = get_paragraph_spacing_points(paragraph)
            if before_pt != COVER_TITLE_SPEC.before_pt or after_pt != COVER_TITLE_SPEC.after_pt:
                failures.append(f'cover title spacing mismatch: {title}')
    return failures


def verify_cover_corner_metadata(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    expected_items = [] if context.expected_cover is None else list(context.expected_cover.get('corner_meta') or [])
    if not expected_items:
        return failures
    cover_paragraphs = get_non_empty_cover_paragraphs(doc)
    by_text = {get_paragraph_text(paragraph): paragraph for paragraph in cover_paragraphs}
    for item in expected_items:
        paragraph = by_text.get(item)
        if paragraph is None:
            failures.append(f'cover corner metadata missing: {item}')
            continue
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            failures.append(f'cover corner metadata not left aligned: {item}')
        if get_run_font_names(paragraph) not in [set(), {COVER_CORNER_SPEC.font_name}]:
            failures.append(f'cover corner font mismatch: {item}')
        if get_run_font_sizes(paragraph) not in [set(), {COVER_CORNER_SPEC.font_size.pt}]:
            failures.append(f'cover corner size mismatch: {item}')
        if get_run_bold_values(paragraph) not in [set(), {False}]:
            failures.append(f'cover corner bold mismatch: {item}')
    return failures


def verify_toc_presence(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    toc_heading, toc_field = find_toc_region(doc)
    if context.expect_toc:
        if toc_heading is None:
            failures.append('TOC heading missing')
        if toc_field is None:
            failures.append('TOC field missing')
        return failures
    if toc_heading is not None:
        failures.append('TOC not expected but heading was retained')
    if toc_field is not None:
        failures.append('TOC not expected but field was retained')
    return failures


def verify_toc_structure(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    toc_heading, toc_field = find_toc_region(doc)
    if not context.expect_toc or toc_heading is None or toc_field is None:
        return failures
    if not is_explicit_toc_heading(get_paragraph_text(toc_heading)):
        failures.append('TOC title text mismatch')
    field_text = get_paragraph_text(toc_field)
    if context.source_kind == 'generated' and context.toc_refresh_state != 'refreshed' and '右键更新目录' not in field_text:
        failures.append('TOC field placeholder missing')
    return failures


def verify_toc_against_body(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    heading_paragraphs = collect_heading_paragraphs(doc)
    expected_titles_by_level: dict[int, list[str]] = {}
    for level, title in context.expected_headings:
        expected_titles_by_level.setdefault(level, []).append(title)
    for level, titles in expected_titles_by_level.items():
        actual_titles = [get_paragraph_text(paragraph) for paragraph in heading_paragraphs.get(level, [])]
        if actual_titles != titles:
            failures.append(f'Heading {level} body/expected mismatch')
    return failures


def verify_structure_preservation(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    if not context.structure_snapshot:
        return failures
    current = snapshot_document_structure(doc)
    for key, before in context.structure_snapshot.items():
        after = current.get(key, 0)
        if after < before:
            failures.append(f'{key} count decreased: before={before} after={after}')
    return failures




def verify_body_paragraph_style(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    for paragraph in collect_body_paragraphs(doc, context):
        text = get_paragraph_text(paragraph)
        if get_run_font_names(paragraph) not in [set(), {BODY_SPEC.font_name}]:
            failures.append(f'body font mismatch: {text}')
        sizes = get_run_font_sizes(paragraph)
        if sizes not in [set(), {BODY_SPEC.font_size.pt}]:
            failures.append(f'body size mismatch: {text}')
        style = paragraph.style
        style_fmt = style.paragraph_format if style is not None else None
        style_left_indent = style_fmt.left_indent.cm if style_fmt is not None and style_fmt.left_indent is not None else 0
        style_first_line_indent = style_fmt.first_line_indent.cm if style_fmt is not None and style_fmt.first_line_indent is not None else 0
        if style_left_indent not in (0, 0.0) or style_first_line_indent not in (0, 0.0):
            failures.append(f'body style indent residue: {text}')
        p_pr = paragraph._p.pPr
        ind = p_pr.ind if p_pr is not None else None
        first_line_chars = ind.get(qn('w:firstLineChars')) if ind is not None else None
        left = ind.get(qn('w:left')) if ind is not None else None
        left_chars = ind.get(qn('w:leftChars')) if ind is not None else None
        if first_line_chars != str(BODY_SPEC.first_line_chars * 100):
            failures.append(f'body first-line indent mismatch: {text}')
        if left not in {None, '0'} or left_chars not in {None, '0'}:
            failures.append(f'body left indent residue: {text}')
        spacing = p_pr.spacing if p_pr is not None else None
        line = spacing.get(qn('w:line')) if spacing is not None else None
        line_rule = spacing.get(qn('w:lineRule')) if spacing is not None else None
        before = spacing.get(qn('w:before')) if spacing is not None else None
        after = spacing.get(qn('w:after')) if spacing is not None else None
        if line != '360' or line_rule != 'auto':
            failures.append(f'body line spacing mismatch: {text}')
        if before not in {None, '0'} or after not in {None, '0'}:
            failures.append(f'body spacing before/after mismatch: {text}')
        if paragraph.alignment != BODY_SPEC.alignment:
            failures.append(f'body alignment mismatch: {text}')
        numpr = extract_paragraph_numpr(paragraph)
        if numpr is not None:
            # bullet list paragraphs have numpr by design — skip them
            num_id_el = numpr.find(qn('w:numId'))
            num_id_val = num_id_el.get(qn('w:val')) if num_id_el is not None else None
            if num_id_val != BULLET_NUMBERING_ID:
                failures.append(f'body numbering residue: {text}')
    return failures


def verify_pagination_rules(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    if context.expect_cover and len(doc.sections) < 2:
        failures.append('cover expected but body section was not split from cover section')
    return failures


def verify_table_text_style(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    for table in collect_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in get_non_empty_paragraphs(cell.paragraphs):
                    if get_run_font_names(paragraph) not in [set(), {TABLE_SPEC.font_name}]:
                        failures.append(f'table font mismatch: {get_paragraph_text(paragraph)}')
                    sizes = get_run_font_sizes(paragraph)
                    if sizes not in [set(), {TABLE_SPEC.font_size.pt}]:
                        failures.append(f'table size mismatch: {get_paragraph_text(paragraph)}')
                    first_line_chars = paragraph._p.pPr.ind.get(qn('w:firstLineChars')) if paragraph._p.pPr is not None and paragraph._p.pPr.ind is not None else None
                    if first_line_chars not in {None, '0'}:
                        failures.append(f'table indent mismatch: {get_paragraph_text(paragraph)}')
    return failures


def verify_table_layout(doc: Document, context: VerificationContext) -> list[str]:
    failures: list[str] = []
    for table in collect_tables(doc):
        if not table.rows:
            failures.append('empty table detected')
            continue
        tbl_w = table._tbl.tblPr.first_child_found_in('w:tblW') if table._tbl.tblPr is not None else None
        if tbl_w is None:
            failures.append('table width missing')
    return failures


_SCRIPT_DIR = Path(__file__).resolve().parent
_DEBUG_XML_SCRIPT = _SCRIPT_DIR / "debug_docx_xml.py"

DEBUG_XML_COMMANDS = (
    f"python3 {_DEBUG_XML_SCRIPT} unpack <input.docx> <output-dir>\n"
    f"python3 {_DEBUG_XML_SCRIPT} validate <output-dir> --original <input.docx>"
)


XML_DEBUG_LANE_HINT_PREFIXES = (
    'toc:',
    'numbering:',
    'toc ',
    'numbering ',
)


XML_DEBUG_LANE_HINT_SUBSTRINGS = (
    'field',
    'section',
    'page-break',
    'structure',
    'image count',
    'table count',
    'relationship',
)


def should_recommend_xml_debug_lane(failures: list[str], context: VerificationContext) -> bool:
    if context.source_kind not in {'generated', 'refreshed', 'legacy'}:
        return False
    for failure in failures:
        if failure.startswith(XML_DEBUG_LANE_HINT_PREFIXES):
            return True
        lowered = failure.lower()
        if any(token in lowered for token in XML_DEBUG_LANE_HINT_SUBSTRINGS):
            return True
    return False


def build_xml_debug_lane_guidance(failures: list[str], context: VerificationContext) -> str:
    if not should_recommend_xml_debug_lane(failures, context):
        return ''
    return (
        "\nThis looks more like a DOCX XML structure or field-binding problem than a normal formatting issue. "
        "Consider switching to the XML debug lane:\n"
        f"{DEBUG_XML_COMMANDS}"
    )


def verify_document_workflow(doc: Document, context: VerificationContext):
    failures = []
    failures.extend(format_failures('cover', verify_cover_presence(doc, context)))
    failures.extend(format_failures('cover', verify_cover_title_block(doc, context)))
    failures.extend(format_failures('cover', verify_cover_corner_metadata(doc, context)))
    failures.extend(format_failures('toc', verify_toc_presence(doc, context)))
    failures.extend(format_failures('toc', verify_toc_structure(doc, context)))
    failures.extend(format_failures('toc', verify_toc_against_body(doc, context)))
    failures.extend(format_failures('body', verify_structure_preservation(doc, context)))
    failures.extend(format_failures('body', verify_body_paragraph_style(doc, context)))
    failures.extend(format_failures('body', verify_pagination_rules(doc, context)))
    failures.extend(format_failures('table', verify_table_text_style(doc, context)))
    failures.extend(format_failures('table', verify_table_layout(doc, context)))
    failures.extend(format_failures('numbering', verify_heading_style_definitions(doc, context)))
    failures.extend(format_failures('numbering', verify_heading_paragraph_instances(doc, context)))
    failures.extend(format_failures('numbering', verify_heading_numbering_definitions(doc, context)))
    failures.extend(format_failures('numbering', verify_heading_numbering_chain(doc, context)))
    if failures:
        guidance = build_xml_debug_lane_guidance(failures, context)
        raise RuntimeError('Verification failed: ' + '; '.join(failures) + guidance)


def verify_heading_conversion(doc: Document, toc_entries: list[tuple[int, str]]):
    context = build_verification_context(
        expected_cover=None,
        expect_cover=False,
        expect_toc=bool(toc_entries),
        expected_headings=toc_entries,
        source_kind='legacy',
        toc_refresh_state='not_refreshed',
    )
    numbering_failures = []
    numbering_failures.extend(verify_heading_style_definitions(doc, context))
    numbering_failures.extend(verify_heading_paragraph_instances(doc, context))
    numbering_failures.extend(verify_heading_numbering_definitions(doc, context))
    numbering_failures.extend(verify_heading_numbering_chain(doc, context))
    toc_failures = verify_toc_against_body(doc, context)
    failures = numbering_failures + toc_failures
    if failures:
        guidance = build_xml_debug_lane_guidance(failures, context)
        raise RuntimeError('Heading verification failed: ' + '; '.join(failures) + guidance)


def strip_heading_number_prefix(text: str) -> str:
    return LEADING_HEADING_NUMBER_RE.sub('', text, count=1).strip()


def is_fence_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith('```') or stripped.startswith('~~~')


def is_short_heading_candidate(text: str) -> bool:
    stripped = text.strip()
    if not stripped or len(stripped) > TXT_HEADING_MAX_LENGTH:
        return False
    if stripped.endswith(('。', '；', '，', '.', ';', ',')):
        return False
    return True


def detect_markdown_heading(line: str, heading_base_level: int) -> tuple[int, str] | None:
    if is_table_line(line):
        return None
    if BULLET_RE.match(line) or ORDERED_RE.match(line):
        return None

    heading = HEADING_RE.match(line)
    if heading:
        source_level = len(heading.group(1))
        level = get_heading_level(source_level, heading_base_level)
        content = normalize_heading_text(heading.group(2).strip())
        if not content:
            return None
        return level, content

    stripped = line.strip()
    if not is_short_heading_candidate(stripped):
        return None

    numbered = MARKDOWN_NUMBERED_HEADING_RE.match(stripped)
    if numbered:
        numbering = numbered.group(1)
        title = numbered.group(2).strip()
        if not title:
            return None
        return min(numbering.count('.') + 1, MAX_HEADING_LEVEL), title

    chinese = MARKDOWN_CHINESE_HEADING_RE.match(stripped)
    if chinese:
        title = chinese.group(2).strip()
        if not title:
            return None
        return 1, title

    return None


def detect_heading_base_level(lines: list[str]) -> int:
    levels = []
    in_code = False
    for line in lines:
        stripped = line.strip()
        if is_fence_line(line):
            in_code = not in_code
            continue
        if in_code:
            continue
        detected = detect_markdown_heading(line, 1)
        if detected:
            levels.append(detected[0])
    return min(levels) if levels else 1


def is_cover_stop_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if is_fence_line(line):
        return True
    if HEADING_RE.match(line) or is_table_line(line):
        return True
    return len(stripped) > 40


def extract_cover_region(lines: list[str]) -> tuple[dict[str, list[str] | str | None], list[str]]:
    cover = {'title': None, 'meta': [], 'corner_meta': []}
    remaining = list(lines)

    used_indices = set()
    has_explicit_cover_labels = False
    for idx, line in enumerate(lines[:8]):
        stripped = line.strip()
        if not stripped:
            continue
        if COVER_CORNER_LABEL_RE.match(stripped):
            cover['corner_meta'].append(stripped)
            used_indices.add(idx)
            continue
        match = COVER_LABEL_RE.match(line)
        if not match:
            continue
        label = (match.group(1) or match.group(2) or '').strip()
        value = match.group(3).strip()
        if not value:
            continue
        has_explicit_cover_labels = True
        used_indices.add(idx)
        if label == '封面标题':
            cover['title'] = value
        else:
            cover['meta'].append(value)

    if has_explicit_cover_labels:
        remaining = [line for idx, line in enumerate(lines) if idx not in used_indices]
        return cover, remaining

    probe: list[tuple[int, str]] = []
    for idx, line in enumerate(lines[:6]):
        if idx in used_indices:
            continue
        stripped = line.strip()
        if not stripped:
            continue
        if is_cover_stop_line(line):
            break
        probe.append((idx, stripped))

    if not probe:
        remaining = [line for idx, line in enumerate(lines) if idx not in used_indices]
        return cover, remaining

    cover['title'] = probe[0][1]
    used_indices.add(probe[0][0])
    for idx, value in probe[1:3]:
        if DATE_RE.search(value) or len(cover['meta']) < 2:
            cover['meta'].append(value)
            used_indices.add(idx)

    remaining = [line for idx, line in enumerate(lines) if idx not in used_indices]
    return cover, remaining


def render_cover(doc: Document, cover: dict[str, list[str] | str | None]):
    title = cover.get('title')
    title_lines = list(cover.get('title_lines') or ([] if not title else [str(title)]))
    meta = cover.get('meta') or []
    corner_meta = cover.get('corner_meta') or []
    for item in corner_meta:
        p = doc.add_paragraph()
        run = p.add_run(str(item))
        set_run_font(run, COVER_CORNER_SPEC.font_name, COVER_CORNER_SPEC.font_size)
        apply_paragraph_style(p, COVER_CORNER_SPEC)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.bold = False
    for _ in range(COVER_TOP_SPACER_COUNT):
        spacer = doc.add_paragraph()
        apply_paragraph_style(spacer, COVER_META_SPEC)
        spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, item in enumerate(title_lines):
        p = doc.add_paragraph()
        run = p.add_run(str(item))
        set_run_font(run, COVER_TITLE_SPEC.font_name, COVER_TITLE_SPEC.font_size)
        apply_paragraph_style(p, COVER_TITLE_SPEC)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.bold = True
        if idx != 0:
            p.paragraph_format.space_before = Pt(0)
    for item in meta:
        p = doc.add_paragraph()
        run = p.add_run(str(item))
        set_run_font(run, COVER_META_SPEC.font_name, COVER_META_SPEC.font_size)
        apply_paragraph_style(p, COVER_META_SPEC)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_explicit_cover(cover_text: str | None) -> dict[str, list[str] | str | None]:
    if cover_text is None:
        return {'title': None, 'title_lines': [], 'meta': [], 'corner_meta': []}

    lines = [line.strip() for line in cover_text.splitlines() if line.strip()]
    if not lines:
        return {'title': None, 'title_lines': [], 'meta': [], 'corner_meta': []}

    return {'title': lines[0], 'title_lines': [lines[0]], 'meta': lines[1:], 'corner_meta': []}


def resolve_cover_decision(
    analysis_cover: dict[str, list[str] | str | None],
    reserve_cover: bool,
    force_cover: bool | None,
    cover_text: str | None,
) -> tuple[dict[str, list[str] | str | None] | None, bool]:
    if force_cover is True:
        return build_explicit_cover(cover_text), True

    if force_cover is False:
        return None, False

    if analysis_cover.get('title') or analysis_cover.get('meta') or analysis_cover.get('corner_meta'):
        return analysis_cover, True

    if reserve_cover:
        return {'title': None, 'meta': [], 'corner_meta': []}, True

    return None, False


def render_cover_placeholder(doc: Document):
    for _ in range(COVER_TOP_SPACER_COUNT):
        spacer = doc.add_paragraph()
        apply_paragraph_style(spacer, COVER_META_SPEC)
        spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder = doc.add_paragraph()
    apply_paragraph_style(placeholder, COVER_META_SPEC)
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER


def render_toc(doc: Document, toc_entries: list[tuple[int, str]]):
    title = doc.add_paragraph()
    title_run = title.add_run(TOC_TITLE)
    set_run_font(title_run, TOC_TITLE_SPEC.font_name, TOC_TITLE_SPEC.font_size)
    apply_paragraph_style(title, TOC_TITLE_SPEC)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_paragraph = doc.add_paragraph()
    apply_paragraph_style(toc_paragraph, TOC_ENTRY_SPEC)
    add_toc_field(toc_paragraph, get_toc_field_levels(toc_entries))

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)


def set_style_outline_level(style, level: int):
    p_pr = style._element.get_or_add_pPr()
    outline = p_pr.get_or_add_outlineLvl()
    outline.set(qn('w:val'), str(min(level, 9) - 1))


def ensure_style_numbering(style, level: int):
    clear_style_numbering(style)

    p_pr = style._element.get_or_add_pPr()
    num_pr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(max(level - 1, 0)))
    num_id = OxmlElement('w:numId')
    num_id.set(qn('w:val'), get_title_numbering_id())
    num_pr.append(ilvl)
    num_pr.append(num_id)
    insert_numpr_after_pstyle(p_pr, num_pr)


def configure_document_styles(doc: Document):
    if 'Normal' in doc.styles:
        normal = doc.styles['Normal']
        normal.font.name = BODY_SPEC.font_name
        normal.font.size = BODY_SPEC.font_size
        r_pr = normal._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement('w:rFonts')
            r_pr.append(r_fonts)
        r_fonts.set(qn('w:ascii'), BODY_SPEC.font_name)
        r_fonts.set(qn('w:hAnsi'), BODY_SPEC.font_name)
        r_fonts.set(qn('w:eastAsia'), BODY_SPEC.font_name)
        apply_style_paragraph_format(normal, BODY_SPEC)

    if '文章信息' in doc.styles:
        cover_corner_style = doc.styles['文章信息']
        cover_corner_style.font.name = COVER_CORNER_SPEC.font_name
        cover_corner_style.font.size = COVER_CORNER_SPEC.font_size
        cover_corner_style.font.bold = False
        apply_style_paragraph_format(cover_corner_style, COVER_CORNER_SPEC)
        r_pr = cover_corner_style._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement('w:rFonts')
            r_pr.append(r_fonts)
        r_fonts.set(qn('w:ascii'), COVER_CORNER_SPEC.font_name)
        r_fonts.set(qn('w:hAnsi'), COVER_CORNER_SPEC.font_name)
        r_fonts.set(qn('w:eastAsia'), COVER_CORNER_SPEC.font_name)

    for level, spec in {
        1: H1_SPEC,
        2: H2_SPEC,
        3: H3_SPEC,
        4: H4_SPEC,
        5: H5_SPEC,
        6: H6_SPEC,
        7: H7_SPEC,
        8: H8_SPEC,
        9: H9_SPEC,
    }.items():
        style = get_or_create_paragraph_style(doc, get_heading_style_name(level), get_heading_style_id(level), spec)
        set_style_outline_level(style, level)
        ensure_style_numbering(style, level)

    ensure_numbering(doc)
    detach_inactive_heading_style_links(doc, get_title_numbering_id())


def configure_document(doc: Document):
    clear_document_body(doc)

    section = doc.sections[0]
    apply_section_page_setup(section)
    configure_document_styles(doc)
    configure_section_footer(section, show_page_number=False, page_number_start=None)


def is_table_line(line: str) -> bool:
    s = line.strip()
    return s.startswith('|') and s.endswith('|') and s.count('|') >= 2


def is_separator_line(line: str) -> bool:
    s = line.strip().strip('|').strip()
    if not s:
        return False
    parts = [p.strip() for p in s.split('|')]
    return all(parts) and all(re.fullmatch(r':?-{3,}:?', p) for p in parts)


def split_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip('|').split('|')]


def set_table_width(table, col_count: int):
    table.autofit = False
    per_col_dxa = max(int(TABLE_CONTENT_WIDTH_DXA / max(col_count, 1)), 1)
    per_col_cm = TABLE_CONTENT_WIDTH_DXA / 567.0 / max(col_count, 1)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(TABLE_CONTENT_WIDTH_DXA))
    tbl_w.set(qn('w:type'), 'dxa')

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells[:col_count]:
            cell.width = Cm(per_col_cm)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in('w:tcW')
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(per_col_dxa))
            tc_w.set(qn('w:type'), 'dxa')


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    apply_paragraph_style(p, TABLE_SPEC)
    set_run_font(run, TABLE_SPEC.font_name, TABLE_SPEC.font_size)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold:
        run.bold = True


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def normalize_toc_heading(text: str) -> str:
    return re.sub(r'\s+', ' ', text.strip().lower())


def is_explicit_toc_heading(text: str) -> bool:
    return normalize_toc_heading(text) in TOC_HEADING_MARKERS


def scan_markdown_toc(lines: list[str]) -> bool:
    in_code = False
    non_empty_seen = 0
    heading_base_level = detect_heading_base_level(lines)
    for line in lines:
        stripped = line.strip()
        if is_fence_line(line):
            in_code = not in_code
            continue
        if in_code or not stripped:
            continue
        non_empty_seen += 1
        detected = detect_markdown_heading(line, heading_base_level)
        content = detected[1] if detected else stripped
        if is_explicit_toc_heading(content):
            return True
        if non_empty_seen >= 20:
            break
    return False


def collect_markdown_toc_entries(lines: list[str]) -> list[tuple[int, str]]:
    entries: list[tuple[int, str]] = []
    heading_base_level = detect_heading_base_level(lines)
    in_code = False
    for line in lines:
        stripped = line.strip()
        if is_fence_line(line):
            in_code = not in_code
            continue
        if in_code:
            continue
        detected = detect_markdown_heading(line, heading_base_level)
        if not detected:
            continue
        level, content = detected
        if not content or is_explicit_toc_heading(content):
            continue
        entries.append((level, content))
    return entries


def detect_txt_heading(line: str) -> tuple[int, str] | None:
    stripped = line.strip()
    if not stripped or len(stripped) > TXT_HEADING_MAX_LENGTH:
        return None
    if stripped.endswith(('。', '；', '，', '.', ';', ',')):
        return None

    numbered = TXT_NUMBERED_HEADING_RE.match(line)
    if numbered:
        numbering = numbered.group(1)
        title = numbered.group(2).strip()
        if not title:
            return None
        level = min(numbering.count('.') + 1, MAX_HEADING_LEVEL)
        return level, title

    chinese = TXT_CHINESE_HEADING_RE.match(line)
    if chinese:
        title = chinese.group(2).strip()
        if not title:
            return None
        return 1, title

    return None


def scan_txt_toc(lines: list[str]) -> bool:
    non_empty_seen = 0
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        non_empty_seen += 1
        if is_explicit_toc_heading(stripped):
            return True
        detected = detect_txt_heading(stripped)
        if detected and is_explicit_toc_heading(detected[1]):
            return True
        if non_empty_seen >= 20:
            break
    return False


def build_txt_blocks(lines: list[str]) -> list[TxtBlock]:
    blocks: list[TxtBlock] = []
    paragraphs: list[str] = []

    def flush_paragraphs():
        nonlocal paragraphs
        if paragraphs:
            blocks.append(TxtBlock('paragraph', paragraphs))
            paragraphs = []

    for line in lines:
        stripped = line.strip()
        detected = detect_txt_heading(line)
        if detected:
            flush_paragraphs()
            level, title = detected
            blocks.append(TxtBlock('heading', [line], level, title))
            continue
        if not stripped:
            flush_paragraphs()
            continue
        paragraphs.append(line)

    flush_paragraphs()
    return blocks


def collect_txt_toc_entries(blocks: list[TxtBlock]) -> list[tuple[int, str]]:
    entries: list[tuple[int, str]] = []
    for block in blocks:
        if block.kind != 'heading' or block.level is None or not block.title:
            continue
        if is_explicit_toc_heading(block.title):
            continue
        entries.append((block.level, block.title))
    return entries


def check_runtime_dependencies(input_path: Path):
    if DOCX_IMPORT_ERROR is not None:
        raise SystemExit(
            'Missing dependency: python-docx. Install it with "python3 -m pip install python-docx" before using this skill.'
        ) from DOCX_IMPORT_ERROR

    if input_path.suffix.lower() not in SUPPORTED_INPUT_SUFFIXES:
        supported = ', '.join(sorted(SUPPORTED_INPUT_SUFFIXES))
        raise SystemExit(f'Unsupported input type: {input_path.suffix.lower()}. Supported types: {supported}')


def analyze_markdown(text: str) -> AnalysisResult:
    raw_lines = text.splitlines()
    cover, body_lines = extract_cover_region(raw_lines)
    return AnalysisResult(cover, body_lines, scan_markdown_toc(body_lines), collect_markdown_toc_entries(body_lines))


def analyze_txt(text: str) -> tuple[AnalysisResult, list[TxtBlock]]:
    raw_lines = text.splitlines()
    cover, body_lines = extract_cover_region(raw_lines)
    blocks = build_txt_blocks(body_lines)
    return AnalysisResult(cover, body_lines, scan_txt_toc(body_lines), collect_txt_toc_entries(blocks)), blocks


def render_markdown(
    doc: Document,
    analysis: AnalysisResult,
    reserve_cover: bool,
    auto_toc: bool,
    force_cover: bool | None = None,
    cover_text: str | None = None,
    force_toc: bool | None = None,
):
    cover, has_cover_section = resolve_cover_decision(analysis.cover, reserve_cover, force_cover, cover_text)
    lines = analysis.body_lines
    if has_cover_section:
        if cover and (cover.get('title') or cover.get('meta') or cover.get('corner_meta')):
            render_cover(doc, cover)
        else:
            render_cover_placeholder(doc)
        configure_section_footer(doc.sections[0], show_page_number=False, page_number_start=None)
        start_body_section(doc)

    should_render_toc = force_toc if force_toc is not None else auto_toc
    if should_render_toc and not analysis.toc_exists and analysis.toc_entries:
        render_toc(doc, analysis.toc_entries)

    heading_base_level = detect_heading_base_level(lines)
    heading_counters = [0] * MAX_HEADING_LEVEL
    in_code = False
    code_lines: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if is_fence_line(line):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph()
                for idx, code_line in enumerate(code_lines):
                    run = p.add_run(code_line)
                    set_run_font(run, CODE_SPEC.font_name, CODE_SPEC.font_size)
                    if idx != len(code_lines) - 1:
                        run.add_break()
                apply_code_block_style(p)
                in_code = False
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if is_table_line(line):
            table_lines = [line]
            j = i + 1
            while j < len(lines) and is_table_line(lines[j]):
                table_lines.append(lines[j])
                j += 1

            rows = [split_table_row(tl) for tl in table_lines]
            has_header_sep = len(rows) >= 2 and is_separator_line(table_lines[1])
            if has_header_sep:
                header = rows[0]
                body = rows[2:]
            else:
                header = None
                body = rows

            col_count = max(len(r) for r in ([header] if header else []) + body)
            table_rows = (1 if header else 0) + max(len(body), 1)
            table = doc.add_table(rows=table_rows, cols=col_count)
            table.style = 'Table Grid'
            set_table_width(table, col_count)
            row_idx = 0
            if header:
                for col, val in enumerate(header):
                    cell = table.rows[0].cells[col]
                    set_cell_text(cell, val, bold=True)
                    shade_cell(cell, TABLE_HEADER_SHADE)
                row_idx = 1
            body_rows = body if body else [[]]
            for row in body_rows:
                for col in range(col_count):
                    val = row[col] if col < len(row) else ''
                    set_cell_text(table.rows[row_idx].cells[col], val, bold=False)
                row_idx += 1
            i = j
            continue

        detected_heading = detect_markdown_heading(line, heading_base_level)
        if detected_heading:
            level, content = detected_heading
            if not content or is_explicit_toc_heading(content):
                i += 1
                continue
            p = doc.add_paragraph(style=get_heading_style_id(level))
            run = p.add_run(content)
            spec = get_heading_spec(level)
            set_run_font(run, spec.font_name, spec.font_size)
            apply_paragraph_style(p, spec)
            apply_heading_numbering(p, level)
            i += 1
            continue

        bullet = BULLET_RE.match(line)
        if bullet:
            indent = bullet.group(1)
            bullet_level = min(len(indent.replace('\t', '    ')) // 2, 2)
            p = doc.add_paragraph()
            run = p.add_run(bullet.group(2).strip())
            set_run_font(run, BODY_SPEC.font_name, BODY_SPEC.font_size)
            apply_bullet_style(p, bullet_level)
            i += 1
            continue

        ordered = ORDERED_RE.match(line)
        if ordered:
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(ordered.group(1).strip())
            set_run_font(run, BODY_SPEC.font_name, BODY_SPEC.font_size)
            apply_paragraph_style(p, BODY_SPEC)
            i += 1
            continue

        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, BODY_SPEC.font_name, BODY_SPEC.font_size)
        apply_paragraph_style(p, BODY_SPEC)
        i += 1

    if in_code and code_lines:
        p = doc.add_paragraph()
        for idx, code_line in enumerate(code_lines):
            run = p.add_run(code_line)
            set_run_font(run, CODE_SPEC.font_name, CODE_SPEC.font_size)
            if idx != len(code_lines) - 1:
                run.add_break()
        apply_code_block_style(p)


def render_txt(
    doc: Document,
    analysis: AnalysisResult,
    blocks: list[TxtBlock],
    reserve_cover: bool,
    auto_toc: bool,
    force_cover: bool | None = None,
    cover_text: str | None = None,
    force_toc: bool | None = None,
):
    cover, has_cover_section = resolve_cover_decision(analysis.cover, reserve_cover, force_cover, cover_text)
    if has_cover_section:
        if cover and (cover.get('title') or cover.get('meta') or cover.get('corner_meta')):
            render_cover(doc, cover)
        else:
            render_cover_placeholder(doc)
        configure_section_footer(doc.sections[0], show_page_number=False, page_number_start=None)
        start_body_section(doc)

    should_render_toc = force_toc if force_toc is not None else auto_toc
    if should_render_toc and not analysis.toc_exists and analysis.toc_entries:
        render_toc(doc, analysis.toc_entries)

    heading_counters = [0] * MAX_HEADING_LEVEL
    for block in blocks:
        if block.kind == 'heading' and block.level is not None and block.title is not None:
            p = doc.add_paragraph(style=get_heading_style_id(block.level))
            run = p.add_run(normalize_heading_text(block.title))
            spec = get_heading_spec(block.level)
            set_run_font(run, spec.font_name, spec.font_size)
            apply_paragraph_style(p, spec)
            apply_heading_numbering(p, block.level)
            continue

        p = doc.add_paragraph()
        for idx, line in enumerate(block.lines):
            run = p.add_run(line)
            set_run_font(run, BODY_SPEC.font_name, BODY_SPEC.font_size)
            if idx != len(block.lines) - 1:
                run.add_break()
        apply_paragraph_style(p, BODY_SPEC)


def convert(
    input_path: Path,
    output_path: Path,
    reserve_cover: bool = False,
    auto_toc: bool = False,
    force_cover: bool | None = None,
    cover_text: str | None = None,
    force_toc: bool | None = None,
):
    suffix = input_path.suffix.lower()
    if suffix not in SUPPORTED_INPUT_SUFFIXES:
        raise ValueError(f'Unsupported input type: {suffix}')

    if suffix == '.docx':
        refresh_existing_docx(
            input_path,
            output_path,
            force_cover=force_cover,
            force_toc=force_toc,
        )
        return

    text = input_path.read_text(encoding='utf-8')
    doc = Document()
    configure_document(doc)

    has_cover_section = False
    cover = None
    if suffix in {'.md', '.markdown'}:
        analysis = analyze_markdown(text)
        cover, has_cover_section = resolve_cover_decision(analysis.cover, reserve_cover, force_cover, cover_text)
        render_markdown(
            doc,
            analysis,
            reserve_cover,
            auto_toc,
            force_cover=force_cover,
            cover_text=cover_text,
            force_toc=force_toc,
        )
    else:
        analysis, blocks = analyze_txt(text)
        cover, has_cover_section = resolve_cover_decision(analysis.cover, reserve_cover, force_cover, cover_text)
        render_txt(
            doc,
            analysis,
            blocks,
            reserve_cover,
            auto_toc,
            force_cover=force_cover,
            cover_text=cover_text,
            force_toc=force_toc,
        )

    if not has_cover_section:
        configure_section_footer(doc.sections[0], show_page_number=True, page_number_start=1)

    should_render_toc = force_toc if force_toc is not None else auto_toc
    context = build_verification_context(
        expected_cover=cover,
        expect_cover=has_cover_section,
        expect_toc=bool(should_render_toc and not analysis.toc_exists and analysis.toc_entries),
        expected_headings=analysis.toc_entries,
        source_kind='generated',
        toc_refresh_state='not_refreshed',
    )
    verify_document_workflow(doc, context)
    doc.save(output_path)


def parse_args():
    parser = argparse.ArgumentParser(description='Convert Markdown, TXT, or DOCX to formatted DOCX.')
    parser.add_argument('input', help='Input .md, .markdown, .txt, or .docx file')
    parser.add_argument('output', nargs='?', help='Optional output .docx path')
    parser.add_argument('--reserve-cover', action='store_true', help='Insert a placeholder cover page when no cover is detected')
    parser.add_argument('--auto-toc', action='store_true', help='Insert a generated TOC page when no explicit TOC heading is detected')
    parser.add_argument('--with-cover', dest='force_cover', action='store_const', const=True, default=None, help='Always generate a cover page and bypass automatic cover detection')
    parser.add_argument('--without-cover', dest='force_cover', action='store_const', const=False, help='Never generate a cover page and bypass automatic cover detection')
    parser.add_argument('--cover-text', help='Explicit cover text. The first non-empty line is used as the title and later lines are rendered as centered metadata')
    parser.add_argument('--with-toc', dest='force_toc', action='store_const', const=True, default=None, help='Always request generated TOC insertion when no explicit TOC heading exists')
    parser.add_argument('--without-toc', dest='force_toc', action='store_const', const=False, help='Never insert a generated TOC page')
    return parser.parse_args()


def main():
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        raise SystemExit(f'Input file not found: {input_path}')

    check_runtime_dependencies(input_path)

    if args.output:
        output_path = Path(args.output).expanduser().resolve()
    else:
        output_path = input_path.with_suffix('.docx')

    if args.cover_text is not None and args.force_cover is False:
        raise SystemExit('--cover-text cannot be used together with --without-cover')

    if args.cover_text is not None and args.force_cover is None:
        force_cover = True
    else:
        force_cover = args.force_cover

    convert(
        input_path,
        output_path,
        reserve_cover=args.reserve_cover,
        auto_toc=args.auto_toc,
        force_cover=force_cover,
        cover_text=args.cover_text,
        force_toc=args.force_toc,
    )
    print(output_path)


if __name__ == '__main__':
    main()
